import docx
import openpyxl
import os, re

filename = "20190311Python教學考古選擇題.docx"
doc = docx.Document(filename)
para = doc.paragraphs
print('段落數量： ', len(para),'\n')
i = 0
is_first_topic = False
is_topic = False
save_path = ".\\imgs\\"

def w_img(blob_data, save_path):
    with open(save_path, "wb") as f:
        f.write(blob_data)

for _ in range(0, len(para)):
    try:
        content = para[_]
        # print([content[0], content[4], content[5]])

        if content.text != "": # 判斷文本是否為空值
            topic = re.sub(r'^\d+', '', content.text)
            option = [content.text[0], content.text[1]]
            if len(topic) > 6:
                if [topic[0], topic[4], topic[5]] == ['(', ')', '：'] or [topic[0], topic[2], topic[3]] == ['（', '）', '：']:
                    i += 1
                    is_first_topic = True
                    is_topic = True
                    print(i, topic)
            # match ... case ... 只支援python 3.10以上
            # 為了相容性則不選擇使用
            if option == ['A', '：']: # 選項A
                is_topic = False
                print(content.text)
            if option == ['B', '：']: # 選項B
                is_topic = False
                print(content.text)
            if option == ['C', '：']: # 選項C
                is_topic = False
                print(content.text)
            if option == ['D', '：']: # 選項D
                is_topic = False
                print(content.text)
            if is_topic and is_first_topic == False:
                print(content.text)
            is_first_topic = False
        img = content._element.xpath(".//pic:pic")
        if img:
            print(img)

        #     img = img[0]
        #     rel = img.xpath('.//a:blip/@r:embed')[0]
        #     image_part = doc.part.related_parts[rel]
        #     img_blob = image_part.image.blob
        #     # print(imagepart.image.blob)
        #     w_img(img_blob, "{}img{}.png".format(save_path, i))

    except IndexError:
        print("IndexERROR")


def get_pictures(word_path, result_path): # 讀取圖片
    """
    图片提取
    :param word_path: word路径
    :return: 
    """
    try:
        doc = docx.Document(word_path)
        dict_rel = doc.part._rels
        for rel in dict_rel:
            print(rel)
            rel = dict_rel[rel]
            if "image" in rel.target_ref:
                if not os.path.exists(result_path):
                    os.makedirs(result_path)
                img_name = re.findall("/(.*)", rel.target_ref)[0]
                word_name = os.path.splitext(word_path)[0]
                if os.sep in word_name:
                    new_name = word_name.split('\\')[-1]
                else:
                    new_name = word_name.split('/')[-1]
                img_name = f'{new_name}-'+'.'+f'{img_name}'
                with open(f'{result_path}/{img_name}', "wb") as f:
                    f.write(rel.target_part.blob)
    except:
        pass

def get_word_pictures(): # 讀取圖片
    global filename
    """
    图片提取
    :param word_path: word路径
    :return: 
    """
    try:
        doc = docx.Document(filename)
        dict_rel = doc.part._rels
        print(dict_rel)
        for index, rel in enumerate(dict_rel):
            rel = dict_rel[rel]
            print(index, rel.target_ref)
            # if "image" in rel.target_ref:
            #     if not os.path.exists(result_path):
            #         os.makedirs(result_path)
            #     img_name = re.findall("/(.*)", rel.target_ref)[0]
            #     word_name = os.path.splitext(word_path)[0]
            #     if os.sep in word_name:
            #         new_name = word_name.split('\\')[-1]
            #     else:
            #         new_name = word_name.split('/')[-1]
            #     img_name = f'{new_name}-'+'.'+f'{img_name}'
            #     with open(f'{result_path}/{img_name}', "wb") as f:
            #         f.write(rel.target_part.blob)
    except:
        pass

# get_word_pictures()
# os.chdir("C:\MyProgram\考題練習")
# spam=os.listdir(os.getcwd())
# for i in spam:
#     get_pictures(str(i),os.getcwd())